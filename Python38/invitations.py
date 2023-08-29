#! python3
# invitations - Sends invitations to people whose names are taken from a text file.

import docx, re
file = open('guests.txt','r')
content = file.readlines()
#names = []
#for i in range(0, 4):
#	regex = re.compile('\n')
#	new = re.sub('', content[i])
#	names.append(new)
#print(names)
doc = docx.Document('guests.docx')
for name in content:
	doc.add_paragraph('It would be a pleasure to have the company of','Brush Script Std' )
	doc.add_paragraph(name,'Fis')
	doc.add_paragraph('at 11010 Memory Lane on the Evening of','Brush Script Std')
	doc.add_paragraph('April 1st','Fis')
	doc.add_paragraph('at 7 o\'clock','Brush Script Std')
	doc.add_page_break()
doc.save('guests3.docx')